from __future__ import annotations

import io
import json
import uuid
import zipfile
from types import SimpleNamespace

import pytest

from agent_smith.app.ports.document_processing import FileProcessingError, ProcessingInput
from agent_smith.app.ports.files import FileRecord
from agent_smith.infra.document_processing import (
    SupportedFileTypeDetector,
    create_processor_registry,
    inspect_image,
)
from agent_smith.infra.document_processing.serialization import build_core_artifacts, parse_jsonl
from agent_smith.workers.agent_worker import AgentWorker
from helpers.files import FakeBlobStore, FakeFileCatalog, FakeFileProcessingStore


def _file(mime_type: str, name: str) -> FileRecord:
    file_id = str(uuid.uuid4())
    return FileRecord(
        id=file_id,
        principal_id=str(uuid.uuid4()),
        original_name=name,
        mime_type=mime_type,
        size_bytes=1,
        object_key=f"files/{file_id}/original",
        status="processing",
    )


@pytest.mark.parametrize(
    ("data", "declared", "name", "expected"),
    [
        (b"hello", "text/plain", "a.txt", "text/plain"),
        (b"# Heading", "text/markdown", "a.md", "text/markdown"),
        (b"name,value\na,1\nb,2\n", "text/csv", "a.csv", "text/csv"),
        (b"%PDF-1.7\n", "application/pdf", "a.pdf", "application/pdf"),
    ],
)
def test_detector_uses_content_with_text_hints(data, declared, name, expected) -> None:
    detected = SupportedFileTypeDetector().detect(
        data=data, declared_mime_type=declared, filename=name
    )
    assert detected.mime_type == expected


def test_detector_prefers_binary_content_over_declared_type_and_extension() -> None:
    detected = SupportedFileTypeDetector().detect(
        data=b"%PDF-1.7\n",
        declared_mime_type="text/plain",
        filename="misleading.txt",
    )

    assert detected.mime_type == "application/pdf"


@pytest.mark.asyncio
async def test_text_csv_and_markdown_golden_fixtures_produce_stable_derivatives() -> None:
    registry = create_processor_registry()
    cases = [
        ("text/plain", "a.txt", b"alpha\nbeta", "paragraph", "alpha\nbeta"),
        ("text/markdown", "a.md", b"# Alpha\n\nBeta", "paragraph", "# Alpha\n\nBeta"),
        ("text/csv", "a.csv", b"name,value\na,1\n", "table", None),
    ]
    for mime_type, name, data, kind, expected_text in cases:
        record = _file(mime_type, name)
        result = await registry.resolve(mime_type).process(
            ProcessingInput(file=record, detected_mime_type=mime_type, data=data)
        )
        artifacts = {item.kind: item for item in build_core_artifacts(result.document)}
        assert set(artifacts) == {"normalized_document", "extracted_text", "chunks"}
        normalized = parse_jsonl(artifacts["normalized_document"].data)
        assert normalized[0]["schemaVersion"] == "1"
        assert normalized[0]["fileId"] == record.id
        assert normalized[1]["kind"] == kind
        assert normalized[1]["text"] == expected_text
        if mime_type == "text/csv":
            assert normalized[1]["table"]["rows"] == [["name", "value"], ["a", "1"]]
            assert normalized[1]["provenance"]["row_start"] == 1
            assert normalized[1]["provenance"]["row_end"] == 2


@pytest.mark.asyncio
async def test_docx_and_xlsx_golden_fixtures_keep_table_boundaries_and_provenance() -> None:
    from docx import Document
    from openpyxl import Workbook

    docx_buffer = io.BytesIO()
    document = Document()
    document.add_heading("Overview", level=1)
    document.add_paragraph("A paragraph")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "42"
    document.save(docx_buffer)

    xlsx_buffer = io.BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Revenue"
    sheet.append(["Month", "Value"])
    sheet.append(["Jan", 42])
    workbook.save(xlsx_buffer)

    registry = create_processor_registry()
    docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    xlsx_mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    docx_result = await registry.resolve(docx_mime).process(
        ProcessingInput(
            file=_file(docx_mime, "a.docx"),
            detected_mime_type=docx_mime,
            data=docx_buffer.getvalue(),
        )
    )
    xlsx_result = await registry.resolve(xlsx_mime).process(
        ProcessingInput(
            file=_file(xlsx_mime, "a.xlsx"),
            detected_mime_type=xlsx_mime,
            data=xlsx_buffer.getvalue(),
        )
    )

    assert [block.kind for block in docx_result.document.blocks] == [
        "heading",
        "paragraph",
        "table",
    ]
    assert xlsx_result.document.blocks[0].provenance.sheet == "Revenue"
    assert xlsx_result.document.blocks[0].provenance.cell_range == "A1:B2"


@pytest.mark.asyncio
async def test_corrupt_docx_and_xlsx_are_non_retryable() -> None:
    registry = create_processor_registry()
    cases = [
        (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "bad.docx",
        ),
        (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "bad.xlsx",
        ),
    ]
    for mime_type, name in cases:
        with pytest.raises(FileProcessingError) as exc:
            await registry.resolve(mime_type).process(
                ProcessingInput(
                    file=_file(mime_type, name),
                    detected_mime_type=mime_type,
                    data=b"PK\x03\x04corrupt",
                )
            )
        assert exc.value.code == "corrupt_document"
        assert exc.value.retryable is False


def test_image_validation_and_corrupt_or_legacy_behavior() -> None:
    from PIL import Image

    buffer = io.BytesIO()
    Image.new("RGB", (3, 2), color="red").save(buffer, format="PNG")
    mime_type, metadata = inspect_image(buffer.getvalue(), declared_mime_type="image/png")
    assert mime_type == "image/png"
    assert metadata["width"] == 3 and metadata["height"] == 2

    with pytest.raises(FileProcessingError) as corrupt:
        inspect_image(b"not an image", declared_mime_type="image/png")
    assert corrupt.value.retryable is False

    with pytest.raises(FileProcessingError) as legacy:
        SupportedFileTypeDetector().detect(
            data=b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1rest",
            declared_mime_type="application/msword",
            filename="a.doc",
        )
    assert legacy.value.code == "unsupported_legacy_format"


def test_chunks_are_jsonl_with_source_provenance() -> None:
    registry = create_processor_registry()
    record = _file("text/plain", "a.txt")

    async def run():
        return await registry.resolve("text/plain").process(
            ProcessingInput(
                file=record,
                detected_mime_type="text/plain",
                data=("needle " * 1000).encode(),
            )
        )

    result = __import__("asyncio").run(run())
    chunks = next(item for item in build_core_artifacts(result.document) if item.kind == "chunks")
    rows = [json.loads(line) for line in chunks.data.decode().splitlines()]
    assert len(rows) > 1
    assert rows[0]["blockIds"] == ["b000001"]


@pytest.mark.asyncio
async def test_pdf_text_layer_and_non_text_pdf_behavior() -> None:
    from pypdf import PdfWriter

    mime_type = "application/pdf"
    registry = create_processor_registry()
    processor = registry.resolve(mime_type)
    result = await processor.process(
        ProcessingInput(
            file=_file(mime_type, "text.pdf"),
            detected_mime_type=mime_type,
            data=_text_pdf("Revenue 42"),
        )
    )
    assert result.document.blocks[0].provenance.page == 1
    assert "Revenue 42" in result.document.blocks[0].text

    scanned = io.BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.write(scanned)
    with pytest.raises(FileProcessingError) as exc:
        await processor.process(
            ProcessingInput(
                file=_file(mime_type, "scan.pdf"),
                detected_mime_type=mime_type,
                data=scanned.getvalue(),
            )
        )
    assert exc.value.code == "scanned_pdf_unsupported"


@pytest.mark.asyncio
async def test_pdf_password_protection_and_corruption_are_non_retryable() -> None:
    from pypdf import PdfReader, PdfWriter

    mime_type = "application/pdf"
    processor = create_processor_registry().resolve(mime_type)
    encrypted = io.BytesIO()
    writer = PdfWriter()
    writer.clone_document_from_reader(PdfReader(io.BytesIO(_text_pdf("secret"))))
    writer.encrypt("password")
    writer.write(encrypted)

    cases = [
        (encrypted.getvalue(), "password_protected"),
        (b"%PDF-1.7\nthis is not a valid PDF", "corrupt_document"),
    ]
    for data, code in cases:
        with pytest.raises(FileProcessingError) as exc:
            await processor.process(
                ProcessingInput(
                    file=_file(mime_type, "broken.pdf"),
                    detected_mime_type=mime_type,
                    data=data,
                )
            )
        assert exc.value.code == code
        assert exc.value.retryable is False


def test_ooxml_sniffing_rejects_fake_zip_and_high_ratio_payload() -> None:
    detector = SupportedFileTypeDetector()
    fake_zip = io.BytesIO()
    with zipfile.ZipFile(fake_zip, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("random.txt", "not OOXML")
    with pytest.raises(FileProcessingError) as unsupported:
        detector.detect(
            data=fake_zip.getvalue(),
            declared_mime_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
            filename="fake.docx",
        )
    assert unsupported.value.code == "unsupported_file_type"

    high_ratio_zip = io.BytesIO()
    with zipfile.ZipFile(high_ratio_zip, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", b"A" * (11 * 1024 * 1024))
    with pytest.raises(FileProcessingError) as bomb:
        detector.detect(
            data=high_ratio_zip.getvalue(),
            declared_mime_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
            filename="bomb.docx",
        )
    assert bomb.value.code == "zip_bomb_detected"
    assert bomb.value.retryable is False


@pytest.mark.asyncio
async def test_worker_processes_uploaded_text_to_persisted_derivatives() -> None:
    catalog, blobs = FakeFileCatalog(), FakeBlobStore()
    store = FakeFileProcessingStore(catalog)
    file_id = str(uuid.uuid4())
    record = FileRecord(
        id=file_id,
        principal_id="principal-a",
        original_name="notes.txt",
        mime_type="text/plain",
        size_bytes=11,
        object_key=f"principals/principal-a/files/{file_id}/original",
        status="pending_upload",
    )
    catalog.records[file_id] = record
    blobs.objects[record.object_key] = (b"hello world", "text/plain", None)
    queued = await store.mark_uploaded_and_enqueue(
        file_id=file_id,
        principal_id="principal-a",
        etag="etag",
        sha256=None,
        pipeline_version="document-v1",
        max_attempts=5,
    )
    assert queued is not None
    container = SimpleNamespace(
        settings=SimpleNamespace(
            file_processing_lease_seconds=60,
            file_processing_timeout_seconds=30,
            file_processing_heartbeat_seconds=20,
            file_processing_poll_seconds=1,
            file_max_bytes=1024,
        ),
        file_processing_store=store,
        blob_store=blobs,
    )
    worker = AgentWorker(container)  # type: ignore[arg-type]

    assert await worker.run_processing_once() is True
    assert catalog.records[file_id].status == "ready"
    assert {item.kind for item in store.derivatives[file_id]} == {
        "normalized_document",
        "extracted_text",
        "chunks",
    }
    assert await worker.run_processing_once() is False


@pytest.mark.asyncio
async def test_worker_classifies_storage_retry_and_corrupt_document_failure() -> None:
    catalog, blobs = FakeFileCatalog(), FakeBlobStore()
    store = FakeFileProcessingStore(catalog)
    settings = SimpleNamespace(
        file_processing_lease_seconds=60,
        file_processing_timeout_seconds=30,
        file_processing_heartbeat_seconds=20,
        file_processing_poll_seconds=1,
        file_max_bytes=1024,
    )

    async def queue(data: bytes, suffix: str) -> FileRecord:
        file_id = str(uuid.uuid4())
        record = FileRecord(
            id=file_id,
            principal_id="principal-a",
            original_name=f"{suffix}.txt",
            mime_type="text/plain",
            size_bytes=len(data),
            object_key=f"principals/principal-a/files/{file_id}/original",
            status="pending_upload",
        )
        catalog.records[file_id] = record
        blobs.objects[record.object_key] = (data, "text/plain", None)
        assert await store.mark_uploaded_and_enqueue(
            file_id=file_id,
            principal_id="principal-a",
            etag="etag",
            sha256=None,
            pipeline_version="document-v1",
            max_attempts=5,
        )
        return record

    retryable = await queue(b"temporary", "retry")
    worker = AgentWorker(
        SimpleNamespace(
            settings=settings,
            file_processing_store=store,
            blob_store=blobs,
        )
    )  # type: ignore[arg-type]
    blobs.fail = True
    assert await worker.run_processing_once() is True
    blobs.fail = False
    assert store.jobs[retryable.id].status == "retry_wait"
    assert catalog.records[retryable.id].status == "processing"
    assert retryable.object_key in blobs.objects

    corrupt = await queue(b"\x00\x00", "corrupt")
    assert await worker.run_processing_once() is True
    assert store.jobs[corrupt.id].status == "failed"
    assert catalog.records[corrupt.id].status == "failed"
    assert catalog.records[corrupt.id].failure_reason == "unsupported_file_type"
    assert corrupt.object_key in blobs.objects


def _text_pdf(text: str) -> bytes:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    parts = [b"%PDF-1.4\n"]
    offsets = [0]

    def add(number: int, body: bytes) -> None:
        offsets.append(sum(len(part) for part in parts))
        parts.append(f"{number} 0 obj\n".encode() + body + b"\nendobj\n")

    add(1, b"<< /Type /Catalog /Pages 2 0 R >>")
    add(2, b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    add(
        3,
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
    )
    add(4, b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode()
    add(5, f"<< /Length {len(stream)} >>\nstream\n".encode() + stream + b"\nendstream")
    xref = sum(len(part) for part in parts)
    parts.append(b"xref\n0 6\n0000000000 65535 f \n")
    parts.extend(f"{offset:010d} 00000 n \n".encode() for offset in offsets[1:])
    parts.append(f"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode())
    return b"".join(parts)
