"""Format-specific processors producing one normalized representation."""

from __future__ import annotations

import asyncio
import csv
import io
import math
import zipfile
from datetime import date, datetime
from typing import Any, Iterable

from agent_smith.app.ports.document_processing import (
    BlockProvenance,
    FileProcessingError,
    FileProcessor,
    NormalizedBlock,
    NormalizedDocument,
    NormalizedTable,
    ProcessingInput,
    ProcessingResult,
)

MAX_EXTRACTED_CHARS = 10_000_000
MAX_TABLE_CELLS = 1_000_000
MAX_PDF_PAGES = 2_000


class ProcessorRegistry:
    def __init__(self, processors: Iterable[FileProcessor]) -> None:
        self._by_mime: dict[str, FileProcessor] = {}
        for processor in processors:
            for mime_type in processor.mime_types:
                if mime_type in self._by_mime:
                    raise ValueError(f"Duplicate processor for {mime_type}")
                self._by_mime[mime_type] = processor

    def resolve(self, mime_type: str) -> FileProcessor:
        processor = self._by_mime.get(mime_type)
        if processor is None:
            raise FileProcessingError(
                "unsupported_file_type", f"No processor supports {mime_type}.", retryable=False
            )
        return processor


class _TextProcessor:
    version = "1"
    mime_types: frozenset[str]

    async def process(self, value: ProcessingInput) -> ProcessingResult:
        text = _decode_text(value.data)
        _check_text_limit(text)
        block = NormalizedBlock(id="b000001", ordinal=0, kind="paragraph", text=text)
        document = NormalizedDocument(
            schema_version="1",
            file_id=value.file.id,
            detected_mime_type=value.detected_mime_type,
            blocks=(block,),
            metadata={"characterCount": len(text)},
        )
        return ProcessingResult(document=document, metadata=document.metadata)


class PlainTextProcessor(_TextProcessor):
    name = "plain_text"
    mime_types = frozenset({"text/plain"})


class MarkdownProcessor(_TextProcessor):
    name = "markdown"
    mime_types = frozenset({"text/markdown"})


class CsvProcessor:
    name = "csv"
    version = "1"
    mime_types = frozenset({"text/csv"})

    async def process(self, value: ProcessingInput) -> ProcessingResult:
        text = _decode_text(value.data)
        _check_text_limit(text)
        try:
            dialect = csv.Sniffer().sniff(text[:8192], delimiters=",;\t")
        except csv.Error:
            dialect = csv.excel
        rows: list[tuple[str, ...]] = []
        cells = 0
        try:
            for row in csv.reader(io.StringIO(text), dialect):
                cells += len(row)
                if cells > MAX_TABLE_CELLS:
                    raise FileProcessingError(
                        "document_limits_exceeded",
                        "CSV contains too many cells.",
                        retryable=False,
                    )
                rows.append(tuple(row))
        except csv.Error as exc:
            raise FileProcessingError(
                "corrupt_document", "CSV content is malformed.", retryable=False
            ) from exc
        table = NormalizedTable(rows=tuple(rows), header_rows=1 if rows else 0)
        block = NormalizedBlock(
            id="b000001",
            ordinal=0,
            kind="table",
            table=table,
            provenance=BlockProvenance(row_start=1, row_end=len(rows)),
        )
        metadata = {"rowCount": len(rows), "cellCount": cells}
        return ProcessingResult(
            document=NormalizedDocument(
                schema_version="1",
                file_id=value.file.id,
                detected_mime_type=value.detected_mime_type,
                blocks=(block,),
                metadata=metadata,
            ),
            metadata=metadata,
        )


class PdfProcessor:
    name = "pypdf_text"
    version = "1"
    mime_types = frozenset({"application/pdf"})

    async def process(self, value: ProcessingInput) -> ProcessingResult:
        return await asyncio.to_thread(self._process_sync, value)

    def _process_sync(self, value: ProcessingInput) -> ProcessingResult:
        try:
            from pypdf import PdfReader
            from pypdf.errors import PdfReadError
        except ImportError as exc:  # pragma: no cover - deployment dependency guard
            raise FileProcessingError(
                "processor_unavailable", "PDF processor is unavailable.", retryable=True
            ) from exc
        try:
            reader = PdfReader(io.BytesIO(value.data), strict=False)
            if reader.is_encrypted:
                raise FileProcessingError(
                    "password_protected", "Password-protected PDF is unsupported.", retryable=False
                )
            if len(reader.pages) > MAX_PDF_PAGES:
                raise FileProcessingError(
                    "document_limits_exceeded", "PDF has too many pages.", retryable=False
                )
            blocks: list[NormalizedBlock] = []
            total = 0
            for index, page in enumerate(reader.pages, start=1):
                text = (page.extract_text() or "").strip()
                if not text:
                    continue
                total += len(text)
                if total > MAX_EXTRACTED_CHARS:
                    raise FileProcessingError(
                        "document_limits_exceeded",
                        "PDF extracted text exceeds safe limits.",
                        retryable=False,
                    )
                blocks.append(
                    NormalizedBlock(
                        id=f"b{len(blocks) + 1:06d}",
                        ordinal=len(blocks),
                        kind="paragraph",
                        text=text,
                        provenance=BlockProvenance(page=index),
                    )
                )
        except FileProcessingError:
            raise
        except (PdfReadError, ValueError, OSError, KeyError) as exc:
            raise FileProcessingError(
                "corrupt_document", "PDF could not be parsed.", retryable=False
            ) from exc
        if not blocks:
            raise FileProcessingError(
                "scanned_pdf_unsupported",
                "PDF has no usable text layer; OCR is not enabled.",
                retryable=False,
            )
        metadata = {"pageCount": len(reader.pages), "characterCount": total}
        return ProcessingResult(
            document=NormalizedDocument(
                schema_version="1",
                file_id=value.file.id,
                detected_mime_type=value.detected_mime_type,
                blocks=tuple(blocks),
                metadata=metadata,
            ),
            metadata=metadata,
        )


class DocxProcessor:
    name = "python_docx"
    version = "1"
    mime_types = frozenset(
        {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    )

    async def process(self, value: ProcessingInput) -> ProcessingResult:
        return await asyncio.to_thread(self._process_sync, value)

    def _process_sync(self, value: ProcessingInput) -> ProcessingResult:
        try:
            from docx import Document
            from docx.opc.exceptions import OpcError
            from docx.table import Table
            from docx.text.paragraph import Paragraph
        except ImportError as exc:  # pragma: no cover
            raise FileProcessingError(
                "processor_unavailable", "DOCX processor is unavailable.", retryable=True
            ) from exc
        try:
            document = Document(io.BytesIO(value.data))
            blocks: list[NormalizedBlock] = []
            section: tuple[str, ...] = ()
            total = 0
            cells = 0
            for item in document.iter_inner_content():
                if isinstance(item, Paragraph):
                    text = item.text.strip()
                    if not text:
                        continue
                    kind = "heading" if (item.style and item.style.name.startswith("Heading")) else "paragraph"
                    if kind == "heading":
                        section = (text,)
                    total += len(text)
                    blocks.append(
                        NormalizedBlock(
                            id=f"b{len(blocks) + 1:06d}",
                            ordinal=len(blocks),
                            kind=kind,
                            text=text,
                            provenance=BlockProvenance(section=section),
                        )
                    )
                elif isinstance(item, Table):
                    rows = tuple(tuple(cell.text.strip() for cell in row.cells) for row in item.rows)
                    cells += sum(len(row) for row in rows)
                    blocks.append(
                        NormalizedBlock(
                            id=f"b{len(blocks) + 1:06d}",
                            ordinal=len(blocks),
                            kind="table",
                            table=NormalizedTable(rows=rows, header_rows=1 if rows else 0),
                            provenance=BlockProvenance(section=section),
                        )
                    )
                if total > MAX_EXTRACTED_CHARS or cells > MAX_TABLE_CELLS:
                    raise FileProcessingError(
                        "document_limits_exceeded", "DOCX exceeds safe limits.", retryable=False
                    )
        except FileProcessingError:
            raise
        except (OpcError, zipfile.BadZipFile, ValueError, KeyError, OSError) as exc:
            raise FileProcessingError(
                "corrupt_document", "DOCX could not be parsed.", retryable=False
            ) from exc
        metadata = {"blockCount": len(blocks), "characterCount": total, "cellCount": cells}
        return ProcessingResult(
            document=NormalizedDocument(
                schema_version="1",
                file_id=value.file.id,
                detected_mime_type=value.detected_mime_type,
                blocks=tuple(blocks),
                metadata=metadata,
            ),
            metadata=metadata,
        )


class XlsxProcessor:
    name = "openpyxl"
    version = "1"
    mime_types = frozenset(
        {"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"}
    )

    async def process(self, value: ProcessingInput) -> ProcessingResult:
        return await asyncio.to_thread(self._process_sync, value)

    def _process_sync(self, value: ProcessingInput) -> ProcessingResult:
        try:
            from openpyxl import load_workbook
            from openpyxl.utils.exceptions import InvalidFileException
            from openpyxl.utils import get_column_letter
        except ImportError as exc:  # pragma: no cover
            raise FileProcessingError(
                "processor_unavailable", "XLSX processor is unavailable.", retryable=True
            ) from exc
        workbook = None
        try:
            workbook = load_workbook(
                io.BytesIO(value.data), read_only=True, data_only=True, keep_links=False
            )
            blocks: list[NormalizedBlock] = []
            cells = 0
            for sheet in workbook.worksheets:
                rows: list[tuple[str, ...]] = []
                max_width = 0
                for raw in sheet.iter_rows(values_only=True):
                    values = tuple(_cell_text(cell) for cell in raw)
                    while values and values[-1] == "":
                        values = values[:-1]
                    if not values and not rows:
                        continue
                    rows.append(values)
                    max_width = max(max_width, len(values))
                    cells += len(values)
                    if cells > MAX_TABLE_CELLS:
                        raise FileProcessingError(
                            "document_limits_exceeded",
                            "XLSX contains too many cells.",
                            retryable=False,
                        )
                if not rows:
                    continue
                cell_range = f"A1:{get_column_letter(max_width)}{len(rows)}" if max_width else None
                blocks.append(
                    NormalizedBlock(
                        id=f"b{len(blocks) + 1:06d}",
                        ordinal=len(blocks),
                        kind="table",
                        table=NormalizedTable(rows=tuple(rows), header_rows=1),
                        provenance=BlockProvenance(
                            sheet=sheet.title,
                            cell_range=cell_range,
                            row_start=1,
                            row_end=len(rows),
                        ),
                    )
                )
        except FileProcessingError:
            raise
        except (
            InvalidFileException,
            zipfile.BadZipFile,
            ValueError,
            KeyError,
            OSError,
            EOFError,
        ) as exc:
            raise FileProcessingError(
                "corrupt_document", "XLSX could not be parsed.", retryable=False
            ) from exc
        finally:
            if workbook is not None:
                workbook.close()
        metadata = {"sheetCount": len(blocks), "cellCount": cells, "formulasRecalculated": False}
        return ProcessingResult(
            document=NormalizedDocument(
                schema_version="1",
                file_id=value.file.id,
                detected_mime_type=value.detected_mime_type,
                blocks=tuple(blocks),
                metadata=metadata,
            ),
            metadata=metadata,
        )


def inspect_image(data: bytes, *, declared_mime_type: str) -> tuple[str, dict[str, Any]]:
    try:
        import warnings

        from PIL import Image, UnidentifiedImageError
    except ImportError as exc:  # pragma: no cover
        raise FileProcessingError(
            "processor_unavailable", "Image validator is unavailable.", retryable=True
        ) from exc
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("error", Image.DecompressionBombWarning)
            with Image.open(io.BytesIO(data)) as image:
                image.verify()
            with Image.open(io.BytesIO(data)) as image:
                image.load()
                fmt = str(image.format or "").upper()
                mime_type = Image.MIME.get(fmt)
                metadata = {
                    "width": image.width,
                    "height": image.height,
                    "format": fmt,
                    "mode": image.mode,
                    "frames": int(getattr(image, "n_frames", 1)),
                }
    except (
        UnidentifiedImageError,
        OSError,
        SyntaxError,
        Image.DecompressionBombError,
        Image.DecompressionBombWarning,
    ) as exc:
        raise FileProcessingError(
            "corrupt_document", "Image could not be decoded safely.", retryable=False
        ) from exc
    if mime_type != declared_mime_type:
        raise FileProcessingError(
            "mime_mismatch", "Image content does not match its declared MIME type.", retryable=False
        )
    return mime_type, metadata


def _decode_text(data: bytes) -> str:
    try:
        return data.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise FileProcessingError(
            "unsupported_encoding", "Text documents must use UTF-8.", retryable=False
        ) from exc


def _check_text_limit(text: str) -> None:
    if len(text) > MAX_EXTRACTED_CHARS:
        raise FileProcessingError(
            "document_limits_exceeded", "Extracted text exceeds safe limits.", retryable=False
        )


def _cell_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.isoformat(sep=" ")
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, float) and math.isfinite(value):
        return str(value)
    return str(value)
